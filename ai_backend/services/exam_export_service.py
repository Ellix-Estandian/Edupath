from docx import Document
from docx.shared import Pt


def export_exam_to_docx(exam: str, output_file: str):
    document = Document()

    title = document.add_heading(
        "EduPath AI Generated Examination",
        level=1,
    )

    title.style.font.size = Pt(20)

    for line in exam.split("\n"):
        document.add_paragraph(line)

    document.save(output_file)